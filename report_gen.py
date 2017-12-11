try:
    from docx import Document
    from docx.shared import Cm
except:
    print('Python-docx module not found. Please install it first')
import datetime
import argparse
import sys
from spgrps import spacegroups
from lxml import etree
try:
    from halo import Halo
except:
    print('For a nice little spinner install the halo Python module.')
    pass

keys_cif = {'_cell_length_a': 'CELLA', '_cell_length_b': 'CELLB',
    '_cell_length_c': 'CELLC', '_cell_angle_alpha': 'ALPHA',
    '_cell_angle_beta': 'BETA', '_cell_angle_gamma': 'GAMMA',
    '_cell_volume': 'VOLUME', '_cell_formula_units_Z': 'ZERR',
    '_exptl_crystal_density_diffrn': 'DENSITY',
    '_exptl_crystal_size_max': 'SIZEMAX',
    '_exptl_crystal_size_mid': 'SIZEMID', '_exptl_crystal_size_min': 'SIZEMIN',
    '_exptl_absorpt_coefficient_mu': 'MU',
    '_exptl_absorpt_correction_type': 'ABSCORR',
    '_exptl_absorpt_correction_T_min': 'TMIN',
    '_exptl_absorpt_correction_T_max': 'TMAX',
    '_diffrn_reflns_theta_min': 'THETAMIN',
    '_diffrn_reflns_theta_max': 'THETAMAX',
    '_diffrn_reflns_number': 'COLLRFL', '_reflns_number_total': 'UNIQUE',
    '_reflns_number_gt': 'SOMETHING',
    '_diffrn_measured_fraction_theta_max': 'COMPLETENESS',
    '_chemical_formula_sum': 'SUM', '_chemical_formula_moiety': 'MOI',
    '_space_group_crystal_system': 'CRYSSYSTEM',
    '_chemical_formula_weight': 'WEIGHT',
    '_refine_ls_number_reflns': 'DATA',
    '_refine_ls_number_parameters': 'PARAM',
    '_refine_ls_number_restraints': 'RESTR',
    '_refine_ls_restrained_S_all': 'GOOF',
    '_refine_ls_R_factor_all': 'R1ALL', '_refine_ls_R_factor_gt': 'R12SIG',
    '_refine_ls_wR_factor_ref': 'WR2ALL', '_refine_ls_wR_factor_gt': 'WR2SIG',
    '_refine_diff_density_max': 'HIGHPEAK',
    '_refine_diff_density_min': 'DEEPHOLE', '_space_group_IT_number': 'SPGRP'}

xslt = etree.parse('MML2OMML.XSL')

def cifReader(cifFile):
    file = open(cifFile,'r').readlines()
    outDict = {}
    outDict['DATE'] = now.strftime("%d.%m.%Y")
    outDict['IDCODE'] = cifFile.split('_')[0]
    for line in range(len(file)):
        try:
            key = file[line].split()[0]
        except IndexError:
            key = ''
        if key in keys_cif.keys():
            if key == '_chemical_formula_sum' and len(file[line].split())==1:
                value = file[line+1].strip()
            else:
                value = ' '.join(file[line].split()[1:])
            outDict[keys_cif[key]] = value
    return outDict

parser = argparse.ArgumentParser(
    description='Generate pretty Strucutre Reports.')
parser.add_argument(
    '-u', '--user', help='name of the synthetic chemist')
parser.add_argument(
    '-i', '--input', help='name of input CIF-file')
parser.add_argument(
    '-s', '--special', help='special details for the refinement',)
parser.add_argument(
    '-d', '--docx', help='path to the template docx-file')
parser.add_argument(
    '-o', '--out', help='name of output .docx-file')
parser.add_argument(
    '-p', '--picture', help='name of the drawing of the molecular structure')
args = parser.parse_args()

now = datetime.datetime.now()
if not args.input:
    cif = input('Please input filename of input CIF-file: ')
    try:
        replacements = cifReader(cif)
    except:
        print('Please supply a valid file in CIF format')
        sys.exit()
else:
    replacements = cifReader(args.input)
if args.special:
    replacements['SPECIAL'] = args.special
else:
    replacements['SPECIAL'] = input('Any special details for the refinement: ')
if args.user:
    replacements['USER'] = args.user
else:
    replacements['USER'] = input(
        'Please input name of the synthetic chemist: ')
if not args.out:
    args.out = input(
        'Please enter the filename of the report to be generated: ')
replacements['PICTURE'] = './image.JPG'

try:
    spinner = Halo(
        text='Please be patient. I am generating your report.',
        spinner='smiley', color='white')
    spinner.start()
except:
    pass

template = Document(args.docx)
spgrp = spacegroups()
replacements['SIZE'] = '{} x {} x {}'.format(
    replacements['SIZEMIN'], replacements['SIZEMID'], replacements['SIZEMAX'])

for paragraph in template.paragraphs:
    tabStops = paragraph.paragraph_format.tab_stops
    tabStops.add_tab_stop(Cm(1.0))
    tabStops.add_tab_stop(Cm(4.0))
    tabStops.add_tab_stop(Cm(8.0))
    tabStops.add_tab_stop(Cm(12.0)) 
    for key in replacements.keys():
        inline = paragraph.runs
        for i in inline:
            if key in i.text:
                if key == 'SPGRP':
                    newLine = i.text.replace('SPGRP', '')
                    i.text = newLine
                    transform = etree.XSLT(xslt)
                    new_dom = transform(etree.fromstring(
                        spgrp.iucrNumberToMathml(replacements[key])))
                    paragraph._element.append(new_dom.getroot())
                elif key == 'PICTURE':
                    i.text = i.text.replace('PICTURE', '')
                    i.add_picture(replacements['PICTURE'])
#                    inline.add_picture('./image.JPG')
                else:
                    newLine = i.text.replace(key, replacements[key])
                    i.text = newLine
                
                
template.save(args.out)

try:
    spinner.stop()
    spinner.succeed('Report generated')
except:
    pass
